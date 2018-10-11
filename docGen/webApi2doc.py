import codecs
import json
import os
import time
import sys

import natsort
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TITLE = 'ShuoGG'
NORMAL_FONT = '宋体'
NORMAL_FONT_SIZE = 12
# 样式查看: https://blog.csdn.net/ibiao/article/details/78595295
TABLE_STYLE = 'Table Grid'

test_json = """
[
	{
		"className":"LockController",
		"methods":[{
			"methodName":"bind",
			"parameters":["sessUserId","mac","id","qrcode"],
			"uri":"lock/bind"
		},{
			"methodName":"addLock",
			"parameters":["sessUserId","lock"],
			"uri":"lock/addLock"
		},{
			"methodName":"addAuth",
			"parameters":["sessUserId","mobile","mac"],
			"uri":"lock/addAuth"
		},{
			"methodName":"removeAuth",
			"parameters":["sessUserId","mobile","mac"],
			"uri":"lock/removeAuth"
		},{
			"methodName":"changeLockStatus",
			"parameters":["sessUserId","mac","id","qrcode","status"],
			"uri":"lock/changeLockStatus"
		},{
			"methodName":"getLockStatus",
			"parameters":["sessUserId","id"],
			"uri":"lock/getLockStatus"
		},{
			"methodName":"sendAuthApply",
			"parameters":["sessUserId","lockId"],
			"uri":"lock/sendAuthApply"
		},{
			"methodName":"getLockInfo",
			"parameters":["sessUserId","lock"],
			"uri":"lock/getLockInfo"
		}]
	},
	{
		"className":"LockLogController",
		"methods":[{
			"methodName":"listAuthLog",
			"parameters":["sessUserId","startTime","endTime"],
			"uri":"lockLog/listAuthLog"
		},{
			"methodName":"listOperLog",
			"parameters":["sessUserId","startTime","endTime"],
			"uri":"lockLog/listOperLog"
		}]
	}
]

"""


def quick_mkdir(name):
    """
    当前目录下建一个文件夹
    :param name: 文件夹名称
    :return: 新建的文件夹的完整路径
    """
    new_directory = os.getcwd() + '\\' + name + "\\"
    if not os.path.exists(new_directory):
        try:
            os.mkdir(os.getcwd() + '\\' + name)
        except Exception as e:
            print(e)
    return new_directory


def get_suffixfiles_fullpath(suffix):
    """
    获取当前目录下所有.xxx文件的路径 (路径列表自然排序)
    :param suffix: 后缀如".sql" ".java"
    :return: list of str
    """
    sql_files = list(filter(lambda x: x.endswith(suffix), os.listdir(os.getcwd())))
    sqlFilesFullPath = list(map(lambda x: os.getcwd() + '\\' + x, sql_files))
    return natsort.natsorted(sqlFilesFullPath)


def get_file_content(file_path):
    """
    读取文件, 暂时只支持utf8和gbk编码的文件, 自动去除BOM
    :param file_path:
    :return: str
    """
    try:
        with open(file_path, encoding='utf-8') as f1:
            raw = f1.read()
            # 去掉BOM
            bom_head = raw.encode(encoding='utf-8')[:3]
            if bom_head == codecs.BOM_UTF8:
                raw = raw.encode(encoding='utf-8')[3:].decode(encoding='utf-8')
            return raw
    except Exception as e:
        with open(file_path, encoding='GBK') as f2:
            return f2.read()


def gen_docx(controllers, doc_name='api.docx'):
    """
    通过json obj生成docx
    :param controllers: list of Controller obj
    :param doc_name: file full path
    :return: bool
    """
    if controllers is None or len(controllers) == 0:
        print('Controller list is None or Empty!')
        return False
    document = Document()
    # 样式设置
    document.styles['Normal'].font.name = NORMAL_FONT
    document.styles['Normal'].font.size = Pt(NORMAL_FONT_SIZE)
    # noinspection PyProtectedMember
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 1.Head生成
    head_runs = document.add_heading(TITLE, 0).runs
    for per_run in head_runs:
        per_run.font.name = NORMAL_FONT
        # noinspection PyProtectedMember
        per_run._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 2.Sql表格的生成
    c_start_num = 1
    for controller, c_index in zip(controllers, range(c_start_num, c_start_num + len(controllers))):
        # 表Title生成
        class_name = controller['className']
        add_one_head(document, str(c_index) + ' ' + class_name[0].upper() + class_name[1:], 16, True)
        m_start_num = 1
        for method, m_index in zip(controller['methods'], range(m_start_num, m_start_num + len(controller['methods']))):
            # 表格标题
            add_one_head(document, str(c_index) + '.' + str(m_index) + ' ' + method['methodName'], 14, True)
            # 表格生成
            row, col = 5, 2
            docx_table = document.add_table(rows=row, cols=col, style=TABLE_STYLE)
            docx_table.cell(0, 0).text = ''
            docx_table.cell(1, 0).text = '接口名称'
            docx_table.cell(2, 0).text = '接口路径'
            docx_table.cell(3, 0).text = '传入参数'
            docx_table.cell(4, 0).text = '返回结果'
            docx_table.cell(0, 1).text = ''
            docx_table.cell(1, 1).text = method['methodName']
            docx_table.cell(2, 1).text = method['uri']
            docx_table.cell(3, 1).text = ', '.join(method['parameters'])
            docx_table.cell(4, 1).text = '{code:200, msg:'', data:}'
            # 参数说明
            add_one_line(document, '传入参数说明:', 12, True)
            add_one_line(document, (":\n".join(method['parameters']) + ':'))
            add_one_line(document, '返回结果说明:', 12, True)
            add_one_line(document, 'code: 200代表成功 400代表失败\nmsg: 返回信息\ndata: null')
            document.add_paragraph("\n")

    document.save(doc_name)
    return True


def add_one_line(document, text, font_size=12, bold=False):
    label = document.add_paragraph()
    label_class_name = label.add_run(text)
    label_class_name.font.size = Pt(font_size)
    label_class_name.bold = bold


def add_one_head(document, text, font_size=14, bold=False):
    label = document.add_heading()
    label_class_name = label.add_run(text)
    label_class_name.font.size = Pt(font_size)
    label_class_name.bold = bold


def run_transfer(json_text):
    global TITLE
    date_str = time.strftime("%Y%m%d", time.localtime())
    dst_folder_path = quick_mkdir('1.new docx')
    name = '接口文档'
    docx_file_name = name + date_str + '.docx'
    TITLE = name
    if gen_docx(json.loads(json_text), dst_folder_path + docx_file_name):
        print(docx_file_name + '转换成功!!')
    else:
        print(docx_file_name + '转换失败!!')


if __name__ == '__main__':
    json_text = sys.argv[1]
    if json_text is None or json_text == '':
        print('No json text, please check the input arg!')
    else:
        run_transfer(json_text)
